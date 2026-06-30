"""
Gerador do Relatório Final — IF1015 ESAIA 2026.1
CodeCompass — Onboarding Inteligente em Codebases Legados
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = r"c:\CIN-UFPE\TASI06-Eng Software com IA\if1015-es-com-ia-2026-1\Relatorio_Final_CodeCompass.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helper functions ──────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(level, text):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name  = "Calibri"
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    return h

def para(text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=False):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.5)
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def table_simple(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # header row
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            set_font(run, bold=True, size=10)
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "BDD7EE")
        tcPr.append(shd)
    # data rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for run in cells[ci].paragraphs[0].runs:
                set_font(run, size=10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"),  "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
#  CAPA
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("IF1015 — ESAIA · Engenharia de Software com IA · 2026.1")
set_font(run, size=11, color=(100, 100, 100))

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CodeCompass")
set_font(run, size=36, bold=True, color=(31, 57, 100))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Onboarding Inteligente em Codebases Legados com IA")
set_font(run, size=16, color=(68, 114, 196))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Equipe CodeCompass")
set_font(run, size=13, bold=True)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Repositório Git: https://github.com/vbmn-cin-ufpe/if1015-es-com-ia-2026-1")
set_font(run, size=10, color=(0, 70, 127))

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

members = [
    "Victor Barros de Miranda Neves  (vbmn@cin.ufpe.br)",
    "Vinicius Henrique Silva  (vhs@cin.ufpe.br)",
    "Alexandre de Souza Cabral  (asc5@cin.ufpe.br)",
    "Arthur Luis de Farias Alves  (alfa@cin.ufpe.br)",
    "Getulio Junqueira de Queiroz Lima  (gjql@cin.ufpe.br)",
    "Carlos Henrique da Silva Frey  (chsf@cin.ufpe.br)",
]
for m in members:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(m)
    set_font(run, size=11)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Recife, 30 de junho de 2026")
set_font(run, size=11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prof. Vinicius Cardoso Garcia — vcg@cin.ufpe.br")
set_font(run, size=10, color=(100, 100, 100))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SUMÁRIO MANUAL
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "Sumário")
toc_items = [
    ("Disclaimer", ""),
    ("1. Introdução", ""),
    ("2. Metodologia", ""),
    ("3. Movimento 1 — Exposição (Alinhar Estratégia)", ""),
    ("4. Movimento 2 — Composição (Desenhar a Solução)", ""),
    ("5. Movimento 3 — Ensaio (Construir e Testar)", ""),
    ("6. Movimento 4 — Ressonância (Medir e Aprender)", ""),
    ("7. Economicidade do Desenvolvimento Assistido por IA", ""),
    ("8. Discussões Técnicas e Estratégicas", ""),
    ("9. Considerações Éticas", ""),
    ("10. Lições Aprendidas e Reflexões Finais", ""),
    ("11. Referências", ""),
    ("12. Apêndices", ""),
]
for item, _ in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5) if item[0].isdigit() or item.startswith("D") else Cm(0)
    run = p.add_run(item)
    set_font(run, size=11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  DISCLAIMER
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "Disclaimer")
para(
    "Este relatório foi elaborado pela equipe CodeCompass para a disciplina IF1015 — Engenharia de "
    "Software com IA (TASI 6), CIn/UFPE, semestre 2026.1, sob orientação do Prof. Vinicius Cardoso "
    "Garcia. O documento segue a estrutura da Metodologia Sinfonia (Garcia & Medeiros, 2025), "
    "compreendendo os quatro movimentos: Exposição, Composição, Ensaio e Ressonância. Todo o "
    "conteúdo reflete o trabalho efetivamente realizado pela equipe ao longo do semestre, "
    "incluindo o uso extensivo de ferramentas de IA generativa (GitHub Copilot / Claude Opus 4 / "
    "Claude Sonnet 4.6) devidamente documentado no Workflow Document (Apêndice A)."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUÇÃO
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "1. Introdução")

heading(2, "1.1 Contextualização do Problema")
para(
    "O onboarding de novos desenvolvedores em codebases de médio e grande porte é uma das "
    "atividades mais custosas da fase de Manutenção e Evolução de Software (SWEBOK KA 5). "
    "Segundo pesquisas de mercado, o tempo médio de ramp-up até produtividade plena varia "
    "de três a seis meses. Durante esse período, o novo desenvolvedor interrompe os "
    "veteranos com perguntas recorrentes (3–5 h/semana por veterano), comete erros por "
    "desconhecimento de decisões arquiteturais não documentadas e produz código que viola "
    "padrões implícitos do time, gerando retrabalho em code review. O conhecimento sobre a "
    "codebase está disperso em código-fonte, mensagens de commit, Pull Requests, wikis "
    "desatualizadas e, principalmente, na memória dos desenvolvedores veteranos — atores "
    "insubstituíveis, porém sobrecarregados."
)
para(
    "O problema se agrava em sistemas legados, onde a documentação formal é escassa ou "
    "obsoleta e as decisões arquiteturais raramente estão explicadas em artefatos formais. "
    "Alta rotatividade gera ciclos viciosos: cada saída leva conhecimento embora e cada "
    "entrada recomeça o processo do zero, elevando o custo marginal de manutenção e o "
    "risco técnico do sistema."
)

heading(2, "1.2 Objetivo Geral e Objetivos Específicos")
para("Objetivo geral:", bold=True)
para(
    "Desenvolver e validar o CodeCompass, um assistente conversacional de onboarding que "
    "utiliza RAG (Retrieval-Augmented Generation) sobre código-fonte e histórico de commits "
    "para reduzir significativamente o tempo de ramp-up de novos desenvolvedores em "
    "codebases legadas."
)
para("Objetivos específicos:", bold=True)
bullet("Indexar repositórios Git com suporte a 15 linguagens de programação, gerando embeddings semânticos via tree-sitter + OpenAI Embeddings API.")
bullet("Disponibilizar interface conversacional (Chat RAG) para perguntas em linguagem natural sobre a codebase.")
bullet("Gerar automaticamente tours guiados dos módulos mais relevantes ranqueados por complexidade ciclomática, churn e acoplamento.")
bullet("Visualizar o grafo de dependências entre módulos e detectar drift arquitetural entre snapshots.")
bullet("Analisar o histórico de commits para explicar decisões técnicas ('Por quê?').")
bullet("Prover métricas de qualidade, análise de hotspots, auditoria de ações, webhooks GitHub e watchlist com notificações por e-mail.")
bullet("Documentar rigorosamente o processo de desenvolvimento assistido por IA com dados de economicidade.")

heading(2, "1.3 Justificativa do Uso de IA Generativa e LLMs")
para(
    "O problema de onboarding é essencialmente um problema de geração de linguagem natural "
    "contextualizada: o novo desenvolvedor precisa de explicações — não apenas de acesso ao "
    "código. Uma solução tradicional (busca por palavras-chave, grep, documentação estática) "
    "recupera trechos de código, mas não os explica, não sintetiza decisões históricas e não "
    "adapta o nível de detalhe à pergunta feita."
)
para(
    "LLMs de grande escala (Claude, GPT-4) possuem capacidade de compreensão e geração de "
    "linguagem natural suficiente para explicar código, sintetizar histórico de commits e "
    "construir narrativas de onboarding. O padrão RAG (Retrieval-Augmented Generation) "
    "supera a limitação de janela de contexto ao injetar apenas os trechos mais relevantes "
    "para cada pergunta, tornando viável responder perguntas sobre repositórios de centenas "
    "de milhares de linhas de código. Ferramentas tradicionais como Sourcegraph Cody focam "
    "em code completion, não em onboarding estruturado; Swimm exige documentação manual "
    "dos veteranos. O CodeCompass preenche esse gap ao automatizar a geração de "
    "walkthroughs e síntese de decisões históricas."
)

heading(2, "1.4 Visão Geral da Abordagem — Metodologia Sinfonia")
para(
    "O projeto foi estruturado segundo a Metodologia Sinfonia (Garcia & Medeiros, 2025), "
    "percorrendo quatro movimentos: (1) Exposição — alinhamento estratégico, personas, "
    "MVP; (2) Composição — arquitetura C4, catálogo de prompts, decisões técnicas; "
    "(3) Ensaio — implementação completa, testes, segurança, CI/CD; "
    "(4) Ressonância — medição, feedback e análise de economicidade."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  2. METODOLOGIA
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "2. Metodologia")

heading(2, "2.1 As Duas Entregas Indissociáveis")
para(
    "O projeto produziu dois artefatos complementares: (a) a Aplicação CodeCompass — sistema "
    "funcional com 27 funcionalidades implementadas, disponível via Docker Compose — e "
    "(b) o Workflow Document — diário de bordo do desenvolvimento assistido por IA, "
    "mantido desde a fase de pré-proposta e atualizado a cada fase, capturando o uso "
    "qualitativo de IA e os dados de economicidade (consumo de tokens, esforço humano real "
    "e contrafactual). Ambos os artefatos foram desenvolvidos em paralelo e são "
    "indissociáveis: a aplicação evidencia o resultado; o Workflow Document evidencia o "
    "aprendizado e o processo."
)

heading(2, "2.2 Aplicação da Metodologia Sinfonia")
para(
    "Os quatro movimentos estruturaram o trabalho de forma sequencial, com os checkpoints "
    "da disciplina como marcos de validação:"
)
bullet("CP1 — Exposição (Aulas 8–13): definição do domínio D8, personas, MVP, métricas de sucesso.")
bullet("CP2 — Composição (Aulas 14–20): implementação das 8 specs iniciais (SPEC-0001 a SPEC-0008), C4 Model, Catálogo de Prompts, ADRs.")
bullet("CP3 — Ensaio (Aulas 21–29): otimização de performance (18.7× embedding), expansão de linguagens, polimento de UX, auditoria de segurança.")
bullet("Apresentação Final — Ressonância (Aula 30+): métricas finais, análise de economicidade consolidada, decisão estratégica.")

heading(2, "2.3 Gestão do Trabalho em Equipe")
para(
    "A equipe de seis membros organizou-se em sprints semanais, com divisão de "
    "responsabilidades por área de domínio (backend, frontend, infraestrutura, documentação). "
    "A comunicação foi conduzida via WhatsApp e sessões síncronas. O repositório Git "
    "(branch dev → main via Pull Requests) serviu como plataforma central de versionamento "
    "e revisão de código. GitHub Copilot (Agent Mode) foi utilizado por todos os membros "
    "para geração e revisão de código, com supervisão humana obrigatória antes de cada merge."
)

heading(2, "2.4 Workflow Document como Documento Vivo")
para(
    "O Workflow Document (WORKFLOW_DOCUMENT.md) foi mantido atualizado a cada fase, "
    "registrando: ferramentas e modelos utilizados, prompts notáveis, decisões tomadas sem "
    "IA, e as três camadas de economicidade (custo real de IA, esforço humano real, "
    "custo contrafactual). O documento completo consta no Apêndice A deste relatório."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  3. MOVIMENTO 1 — EXPOSIÇÃO
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "3. Movimento 1 — Exposição (Alinhar Estratégia)")
para("Artefatos: Canvas de Estratégia e Ação, Personas, Declaração de Missão/Visão, "
     "Métricas de Sucesso, Matriz de Priorização (Impacto × Esforço), Escopo do MVP.", bold=False)
doc.add_paragraph()

heading(2, "3.1 Canvas de Estratégia e Ação")
para(
    "Problema central: O onboarding de novos desenvolvedores em codebases legadas é lento, "
    "custoso e depende excessivamente de conhecimento tácito de veteranos. "
    "Contexto de negócio: Empresas de tecnologia com alta rotatividade sofrem perda "
    "contínua de conhecimento institucional, elevando o custo marginal de manutenção. "
    "Objetivo de alto nível: Reduzir o tempo de onboarding de 3–6 meses para semanas, "
    "automatizando a geração de documentação contextualizada e respostas a perguntas "
    "sobre o código."
)

heading(2, "3.2 Personas Principais")
para("Persona 1 — Desenvolvedor Novato (usuário primário):", bold=True)
bullet("Perfil: recém-contratado, 1–3 anos de experiência, não conhece a codebase.")
bullet("Dores: interrompe veteranos com perguntas básicas; demora semanas para entender fluxos-chave; comete erros por desconhecimento de padrões implícitos.")
bullet("Ganhos esperados: respostas instantâneas a perguntas sobre o código; tour guiado dos módulos críticos; grafo visual de dependências.")

para("Persona 2 — Tech Lead / Admin (usuário secundário):", bold=True)
bullet("Perfil: 5+ anos de experiência, responsável pela produtividade do time.")
bullet("Dores: interrompido 3–5 h/semana por perguntas de novatos; documentação desatualizada; dificuldade de medir progresso de onboarding.")
bullet("Ganhos esperados: painel de métricas de onboarding; alertas de drift arquitetural; audit log de ações; webhooks para re-indexação automática.")

heading(2, "3.3 Declaração de Missão e Visão")
para(
    "Missão: Democratizar o acesso ao conhecimento de codebases legadas, tornando o "
    "onboarding de novos desenvolvedores mais rápido, autônomo e seguro através de "
    "IA conversacional contextualizada pelo código real."
)
para(
    "Visão: Tornar-se a ferramenta padrão de onboarding técnico para times de engenharia "
    "que mantêm sistemas legados, eliminando a dependência do conhecimento tácito de "
    "veteranos como única fonte de verdade sobre a arquitetura."
)
para(
    "Alinhamento ético: O sistema é transparente sobre suas fontes (RAG cita os chunks "
    "usados), declara incerteza quando o contexto é insuficiente, e não substitui a "
    "revisão humana do código — é uma ferramenta de apoio, não de automação de decisões."
)

heading(2, "3.4 Métricas de Sucesso")
table_simple(
    ["Métrica", "Tipo", "Meta", "Resultado"],
    [
        ["Indexação funcional (15 linguagens)", "Técnica", "✅ Implementado", "✅ Atingido"],
        ["Tempo de resposta no chat", "Técnica", "< 15 segundos", "✅ 3–8s típico"],
        ["Velocidade de embedding (2800 chunks)", "Técnica", "< 30s", "✅ 11.8s (18.7×)"],
        ["Funcionalidades implementadas (MVP)", "Produto", "8 specs MVP", "✅ 27 funcs. totais"],
        ["Testes (unit + integração + E2E)", "Qualidade", "Cobertura das specs core", "✅ Implementados"],
        ["Workflow Document completo", "Processo", "Todas as fases documentadas", "✅ Completo"],
    ],
    [5, 3, 3.5, 3.5],
)

heading(2, "3.5 Matriz de Impacto × Esforço e Priorização")
para(
    "A priorização das funcionalidades seguiu uma Matriz de Impacto × Esforço, classificando "
    "cada feature em quatro quadrantes: Quick Wins (alto impacto, baixo esforço), "
    "Projetos Estratégicos (alto impacto, alto esforço), Fill-ins (baixo impacto, baixo esforço) "
    "e Desperdícios (baixo impacto, alto esforço)."
)
bullet("Quick Wins (Fase 1): Chat RAG, Tour Guiado, Busca Semântica, Autenticação.")
bullet("Projetos Estratégicos (Fase 2–4): Grafo de Dependências, Drift Arquitetural, Webhooks HMAC, Watchlist.")
bullet("Fora do escopo: Integração Slack/Teams, personalização de tour por role, análise profunda de PRs.")

heading(2, "3.6 Escopo do MVP")
para("Features no MVP (implementadas):", bold=True)
bullet("Indexação de repositório (URL remota ou path local) com 15 linguagens")
bullet("Chat RAG conversacional em linguagem natural")
bullet("Tour guiado automático ranqueado por complexidade × churn × acoplamento")
bullet("Grafo de dependências interativo")
bullet("Análise de impacto de mudanças")
bullet("Busca semântica sobre código indexado")
bullet("Histórico de commits com explicações 'Por quê?' via LLM")
bullet("Métricas de qualidade e relatório via LLM")
bullet("Autenticação JWT + verificação de e-mail + reset de senha")
bullet("Observabilidade: logging estruturado, correlation ID, liveness/readiness")

para("Fora do escopo (explicitamente não implementado):", bold=True)
bullet("Integração com Slack/Teams")
bullet("Personalização do tour por role (backend vs frontend dev)")
bullet("Análise profunda de Pull Requests (apenas commits)")
bullet("Deploy on-premise com documentação DevOps completa")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  4. MOVIMENTO 2 — COMPOSIÇÃO
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "4. Movimento 2 — Composição (Desenhar a Solução)")
para("Artefatos: C4 Model (N1–N3), Catálogo de Registros de Prompt, Canvas de Design "
     "de Experimento, Registro de Decisões Arquiteturais (ADRs).")
doc.add_paragraph()

heading(2, "4.1 Arquitetura com C4 Model")

heading(3, "Nível 1 — Contexto")
para(
    "O CodeCompass interage com dois perfis de usuário (Desenvolvedor Novato e Tech Lead/Admin) "
    "e quatro sistemas externos: GitHub/GitLab/Bitbucket (fonte dos repositórios e receptor de "
    "push events via HMAC-SHA256); LLM Provider configurável (Abacus AI / Anthropic / OpenAI) "
    "para geração de linguagem natural; OpenAI Embeddings API (text-embedding-3-small) para "
    "vetorização semântica; e E-mail Gateway (SMTP) para notificações da watchlist. "
    "O LLM é o componente central — toda resposta do chat, geração de tours, relatórios de "
    "qualidade e interpretação de drift arquitetural passam por ele. O sistema opera em modo "
    "fallback (template-based) quando a API do LLM não está disponível."
)

heading(3, "Nível 2 — Contêineres")
table_simple(
    ["Contêiner", "Tecnologia", "Porta", "Responsabilidade"],
    [
        ["Frontend SPA", "React 18 + TypeScript 5.8 + Vite 6 + Tailwind CSS", ":5173",
         "Interface com 10 abas: Chat RAG, Tour, Grafo, Drift, Histórico, Métricas, Repositório, Watchlist, Admin. Dark mode + sidebar colapsável."],
        ["Backend API", "Python 3.11 + FastAPI 0.115 — Arquitetura Hexagonal", ":8000",
         "11 routers REST; pipeline RAG; drift arquitetural; audit middleware; webhooks HMAC-SHA256; watchlist + notificações."],
        ["PostgreSQL 16", "PostgreSQL", ":5432",
         "Usuários, sessões, repos, commits, métricas, audit_log, webhooks, watchlist, snapshots do grafo."],
        ["ChromaDB 0.5", "ChromaDB", ":8001",
         "Vector store para embeddings de código; busca semântica por cosine similarity (RAG)."],
    ],
    [3.5, 5, 1.5, 7],
)
para(
    "Todos os 4 contêineres são orquestrados pelo docker-compose.yml na raiz do repositório. "
    "O frontend consome apenas o backend — sem acesso direto a banco ou vector store."
)

heading(3, "Nível 3 — Componentes do Backend")
para(
    "O Backend API segue Arquitetura Hexagonal (Ports & Adapters) com três camadas: "
    "(a) Controllers — 11 routers FastAPI, responsáveis apenas por orquestração HTTP; "
    "(b) Services — 13 serviços de domínio contendo a lógica de negócio isolada de "
    "frameworks externos; (c) Infrastructure — adaptadores que implementam as Ports "
    "(PostgresAdapter, ChromaAdapter, LlmClient, GitClient, AuditRepository, "
    "WatchlistRepository, WebhookRepository). A injeção de dependência é gerenciada "
    "centralmente em dependencies.py via FastAPI.Depends."
)
para("Serviços de domínio principais:", bold=True)
bullet("repo_service: orquestra git clone → chunking (tree-sitter) → embedding → ChromaDB → PostgreSQL.")
bullet("chat_service: RAG pipeline — embed(question) → ChromaDB.query(top-5) → LLM.generate.")
bullet("tour_service: score = complexidade_ciclomática × churn × acoplamento; walkthroughs via LLM.")
bullet("architecture_drift_service: compara snapshots do grafo → drift_score = changed_elements / total × 100.")
bullet("notification_service: detecta módulos alterados pós-indexação → e-mail para watchlist subscribers.")

heading(2, "4.2 Registro de Decisões Arquiteturais (ADRs)")
table_simple(
    ["ADR", "Decisão", "Trade-off / Justificativa"],
    [
        ["ADR-001", "Arquitetura Hexagonal (Ports & Adapters)",
         "Testabilidade e substituição de adapters sem alterar lógica de negócio. Custo: mais arquivos e indirection."],
        ["ADR-002", "ChromaDB como Vector Store",
         "Open source, roda local sem dependência de cloud. Custo: não escala para bilhões de vetores."],
        ["ADR-003", "FastAPI + BackgroundTasks para indexação",
         "Indexação assíncrona: retorna 202 imediatamente, cliente faz polling. Alternativa Celery seria mais robusta mas mais complexa."],
        ["ADR-004", "OpenAI text-embedding-3-small + ThreadPoolExecutor",
         "18.7× mais rápido que sentence-transformers local (11.8s vs 220s). Custo: dependência de API paga (~$0.00006 por indexação)."],
        ["ADR-005", "Multi-provider LLM via env var",
         "Suporte a Abacus AI, OpenAI e Anthropic sem mudança de código. Flexibilidade de custo e performance."],
        ["ADR-006", "JWT stateless para autenticação",
         "Sem necessidade de session store distribuído. Custo: impossibilidade de revogar tokens antes da expiração."],
        ["ADR-007", "Webhooks GitHub com HMAC-SHA256",
         "Verificação criptográfica da origem do evento. Custo: necessidade de gestão de segredos por webhook."],
        ["ADR-008", "Audit log via middleware FastAPI",
         "Toda mutação (POST/PATCH/DELETE) é registrada automaticamente, sem instrumentação explícita em cada endpoint."],
    ],
    [1.5, 5, 10],
)

heading(2, "4.3 Catálogo de Registros de Prompt")
para("Os seis prompts centrais da solução estão documentados no CATALOGO_PROMPTS.md (Apêndice C). "
     "Abaixo, os três principais:")
doc.add_paragraph()

heading(3, "PROMPT-001 — Assistente RAG de Onboarding (v1.2)")
para("Objetivo: Responder perguntas em linguagem natural sobre a codebase indexada.")
para("Template (simplificado):")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.0)
p.style = doc.styles["No Spacing"]
run = p.add_run(
    "[SYSTEM] Você é um assistente especializado em onboarding de desenvolvedores. "
    "Baseie suas respostas SOMENTE no contexto fornecido...\n"
    "[USER] Com base no código recuperado: {question}\nContexto: {context_chunks}"
)
set_font(run, name="Courier New", size=9)
doc.add_paragraph()
bullet("Taxa de sucesso estimada: 78% em testes manuais (20 perguntas sobre o próprio CodeCompass).")
bullet("Mitigação de alucinações: instrução explícita para declarar insuficiência de contexto; máximo 5 chunks por chamada.")
bullet("Versão 1.2: adicionada restrição 'SOMENTE no contexto fornecido' após observar alucinações na v1.0.")

heading(3, "PROMPT-002 — Explicação 'Por Quê?' via Histórico de Commits (v1.0)")
para("Objetivo: Responder sobre decisões históricas de um módulo a partir de commits classificados.")
bullet("Sintetiza evidências de commits reais em narrativa coerente.")
bullet("Taxa de sucesso: 72% — degrada em repos com mensagens de commit pobres (ex: 'fix', 'update').")

heading(3, "PROMPT-003 — Geração de Tour Guiado (v1.1)")
para("Objetivo: Gerar walkthroughs explicativos para os módulos mais importantes ranqueados por score.")
bullet("Score = complexidade_ciclomática × churn × acoplamento.")
bullet("Instrução para usar linguagem acessível a desenvolvedor iniciante.")

heading(2, "4.4 Canvas de Design de Experimento")
para("Hipótese principal:", bold=True)
para(
    '"Um assistente conversacional alimentado por RAG sobre código-fonte e histórico de commits '
    'consegue reduzir significativamente o tempo que um desenvolvedor novo leva para responder '
    'perguntas básicas sobre uma codebase desconhecida, comparado a ler o código diretamente."'
)
bullet("Métricas: tempo médio de resposta a 10 perguntas padronizadas; taxa de acerto (resposta correta/útil); satisfação qualitativa (escala 1–5).")
bullet("Condições de teste: Grupo A — Dev novato com acesso ao CodeCompass; Grupo B — Dev novato com acesso apenas ao código e documentação estática.")
bullet("Risco principal: alucinações do LLM comprometendo a confiança nas respostas.")
bullet("Critério de decisão: perseverar se taxa de acerto > 70% e satisfação média > 3.5/5.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  5. MOVIMENTO 3 — ENSAIO
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "5. Movimento 3 — Ensaio (Construir e Testar)")
para("Artefatos: Canvas de Testes e Validação, Checklist de Lançamento, evidências de "
     "versionamento e CI/CD.")
doc.add_paragraph()

heading(2, "5.1 Estratégia de Desenvolvimento e Tecnologias")
table_simple(
    ["Camada", "Tecnologia", "Versão", "Papel"],
    [
        ["Backend", "Python + FastAPI", "3.11 / 0.115", "API REST, pipeline RAG, serviços de domínio"],
        ["Frontend", "React + TypeScript + Vite", "18 / 5.8 / 6", "SPA com 10 abas, dark mode, sidebar colapsável"],
        ["Vector Store", "ChromaDB", "0.5", "Embeddings de código para busca semântica"],
        ["Banco de dados", "PostgreSQL", "16", "Persistência relacional completa"],
        ["Embeddings", "OpenAI text-embedding-3-small", "-", "Vetorização semântica (18.7× mais rápido)"],
        ["LLM", "Abacus AI / Anthropic / OpenAI", "-", "Geração de linguagem natural (multi-provider)"],
        ["AST Parsing", "tree-sitter", "-", "Chunking semântico em 15 linguagens"],
        ["Métricas de código", "radon", "-", "Complexidade ciclomática"],
        ["Git", "GitPython", "-", "Clone, log de commits, análise de churn"],
        ["Infra", "Docker + Docker Compose", "-", "Orquestração dos 4 contêineres"],
        ["IaC", "Terraform", "-", "Infraestrutura AWS (ECS Fargate, RDS, ALB)"],
        ["Segurança", "bcrypt + JWT + HMAC-SHA256", "-", "Hashing de senha, tokens e webhooks"],
    ],
    [3, 4, 2.5, 7.5],
)

heading(2, "5.2 Fluxo de Integração com LLMs")
para(
    "O LlmClient (infrastructure/llm_client.py) abstrai os três providers via interface "
    "unificada. O provider é selecionado em runtime via LLM_PROVIDER env var, sem mudança "
    "de código. O sistema implementa fallback template-based quando LLM_API_KEY não está "
    "configurada, garantindo funcionamento básico sem dependência de API paga."
)
para("Salvaguardas implementadas:", bold=True)
bullet("Instrução explícita no SYSTEM prompt para não alucinar quando contexto for insuficiente.")
bullet("Limite de 5 chunks por chamada de chat (controle de custo e qualidade de contexto).")
bullet("Timeout e tratamento de erro robusto em todas as chamadas de API externas.")
bullet("Credenciais gerenciadas via Pydantic BaseSettings com variáveis de ambiente obrigatórias.")

heading(2, "5.3 Canvas de Testes e Validação")

heading(3, "Testes Funcionais")
bullet("Unitários: cobertura dos serviços de domínio (auth_service, chat_service, tour_service, embedding_service, drift_service) com mocks dos adapters de infraestrutura.")
bullet("Integração: testes com PostgreSQL e ChromaDB reais via Docker, validando fluxos completos de indexação → chat → tour.")
bullet("E2E: testes de fluxo completo usuário → frontend → backend → LLM (mock) → resposta.")

heading(3, "Qualidade de Outputs do LLM")
bullet("Testes manuais com 20 perguntas padronizadas sobre o próprio CodeCompass (dogfooding): 78% de utilidade avaliada.")
bullet("Avaliação da capacidade de detectar insuficiência de contexto (instrução 'SOMENTE no contexto fornecido').")
bullet("Edge cases testados: repositório vazio, arquivo binário, commit sem mensagem, pergunta fora do escopo.")

heading(3, "Avaliação de Alucinação e Mitigações")
bullet("Alucinação observada: nas versões iniciais do PROMPT-001 (v1.0), o LLM respondia perguntas além do contexto RAG.")
bullet("Mitigação v1.2: adição de restrição explícita + instrução para declarar insuficiência.")
bullet("Monitoramento contínuo: sistema de feedback (👍/👎) por resposta no chat, com painel admin.")

heading(3, "Performance e Latência")
bullet("Tempo de indexação (2825 chunks, nestjs/nest): 11.8s com OpenAI (vs 220s local — 18.7× mais rápido).")
bullet("Latência de resposta do chat: 3–8s típico (excluindo indexação).")
bullet("Custo de indexação completa do nestjs/nest: ~$0.00006 (praticamente zero).")

heading(2, "5.4 Análise de Segurança (OWASP Top 10)")
table_simple(
    ["Vulnerabilidade", "Status", "Mitigação Implementada"],
    [
        ["A01 — Broken Access Control", "✅ Mitigado", "JWT obrigatório em todos os endpoints protegidos; middleware require_auth; RBAC admin vs user."],
        ["A02 — Cryptographic Failures", "✅ Mitigado", "bcrypt para hash de senhas; HMAC-SHA256 para webhooks; HTTPS para todas as APIs externas."],
        ["A03 — Injection", "✅ Mitigado", "Parâmetros PostgreSQL via psycopg2 (parameterized queries); sem SQL concatenado."],
        ["A04 — Insecure Design", "✅ Mitigado", "Credenciais via env vars obrigatórias (Pydantic BaseSettings); sem segredos hardcoded."],
        ["A05 — Security Misconfiguration", "✅ Mitigado", "CORS configurado explicitamente; headers de segurança; admin seeded via env vars."],
        ["A06 — Vulnerable Components", "✅ Monitorado", "Dependências gerenciadas via pyproject.toml com versões fixas."],
        ["A09 — Security Logging", "✅ Implementado", "Audit log automático via middleware (todas as mutações POST/PATCH/DELETE)."],
        ["Prompt Injection", "✅ Mitigado (parcial)", "Contexto RAG é código indexado — não input direto do usuário. Sem execução de código gerado pelo LLM."],
    ],
    [4, 2, 11],
)

heading(2, "5.5 Evidências de Funcionamento")
para(
    "O sistema está disponível via Docker Compose com comando único (docker compose up --build). "
    "A documentação completa de execução está em COMO_RODAR.md. "
    "Funcionalidades demonstradas: indexação de repositórios públicos (nestjs/nest, "
    "o próprio CodeCompass), chat RAG, tour guiado, grafo de dependências, "
    "drift arquitetural, histórico de commits, métricas, watchlist, webhooks e audit log."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  6. MOVIMENTO 4 — RESSONÂNCIA
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "6. Movimento 4 — Ressonância (Medir e Aprender)")
para("Artefatos: Painel de Feedback e Insights, Canvas de Escalabilidade.")
doc.add_paragraph()

heading(2, "6.1 Lançamento e Coleta de Feedback")
para(
    "O feedback foi coletado por dois métodos: (a) uso simulado por membros da equipe "
    "seguindo as personas definidas na Exposição (desenvolvedor novato indexando o próprio "
    "repositório CodeCompass e fazendo perguntas padronizadas); (b) sistema de feedback "
    "in-app (👍/👎) implementado no chat, com painel admin para análise de qualidade das "
    "respostas do LLM. A avaliação de experiência do usuário foi conduzida por 4 testadores "
    "internos com perguntas sobre o repositório nestjs/nest."
)

heading(2, "6.2 Painel de Feedback e Insights")
para("Resultados quantitativos (avaliação interna — 20 perguntas padronizadas):", bold=True)
bullet("Taxa de respostas úteis/corretas: 78% (PROMPT-001 v1.2).")
bullet("Taxa de respostas com alucinação detectada: 8% (reduzida de ~22% na v1.0).")
bullet("Latência média de resposta: 5.2s.")
bullet("Satisfação geral (escala 1–5): 4.1/5.")

para("Insights qualitativos:", bold=True)
bullet("O tour guiado foi avaliado como 'muito útil' — ranqueamento automático por score poupa horas de leitura do código.")
bullet("O grafo de dependências é especialmente valioso para entender o impacto de mudanças.")
bullet("Perguntas sobre decisões de design ('Por que PostgreSQL em vez de MongoDB?') falham quando a decisão não está documentada em commits — limitação estrutural do RAG.")
bullet("Dark mode e sidebar colapsável melhoram significativamente a experiência em sessões longas.")

heading(2, "6.3 Validação das Hipóteses")
para(
    "Hipótese principal validada: O CodeCompass reduz o tempo de resposta a perguntas "
    "sobre a codebase de minutos (leitura direta do código) para segundos (3–8s). "
    "78% das perguntas padronizadas receberam respostas corretas/úteis, superando o "
    "critério de sucesso de 70%. A hipótese secundária (satisfação > 3.5/5) também foi "
    "confirmada (4.1/5)."
)
para(
    "Hipótese refutada: O tour guiado cobre 100% dos módulos relevantes. Na prática, "
    "o ranqueamento por score favoreça módulos de alta complexidade, podendo subestimar "
    "módulos de negócio críticos mas simples — identificado como ponto de melhoria."
)

heading(2, "6.4 Decisão Estratégica")
para("Decisão: PERSEVERAR — com ajustes de otimização.")
bullet("A validação das hipóteses confirma a proposta de valor central.")
bullet("Ajustes prioritários: melhorar qualidade das respostas para perguntas sobre decisões de design (RAG sobre PRs e wikis); personalização do tour por role.")
bullet("Expansão planejada no Canvas de Escalabilidade.")

heading(2, "6.5 Canvas de Escalabilidade")
para("Caminhos para evolução além do MVP:", bold=True)
bullet("RAG sobre Pull Requests: indexar descrições e comentários de PRs para explicar decisões de design.")
bullet("Personalização de tour por role: frontend dev vs backend dev vs DevOps.")
bullet("Integração Slack/Teams: notificações de mudanças arquiteturais diretamente no canal do time.")
bullet("Análise de sentimento em commits: detectar períodos de estresse técnico (muitos bugfixes seguidos).")
bullet("SaaS multi-tenant: escalabilidade horizontal com Kubernetes, RDS gerenciado, CDN para o frontend.")
bullet("Suporte a monorepos: detecção automática de múltiplos serviços em um único repositório.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  7. ECONOMICIDADE
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "7. Economicidade do Desenvolvimento Assistido por IA")
para(
    "Esta seção consolida as três camadas de economicidade registradas no Workflow Document "
    "ao longo do projeto. O objetivo é responder: qual foi o custo real do desenvolvimento "
    "assistido por IA, comparado a uma estimativa do custo de um desenvolvimento equivalente "
    "feito integralmente por humanos? Cotação USD → BRL utilizada: R$5,50 (junho/2026)."
)

heading(2, "7.1 Camada 1 — Custo Real de IA (Total do Projeto)")
table_simple(
    ["Fase", "Tokens Entrada (est.)", "Tokens Saída (est.)", "Custo IA (USD)", "Custo IA (R$)"],
    [
        ["Pré-Proposta / Exposição", "~30.000", "~57.000", "~$1,08", "~R$5,94"],
        ["Composição (8 specs)", "~186.000", "~266.000", "~$22,74", "~R$125,07"],
        ["Ensaio — Otimização", "~88.000", "~65.000", "~$1,11", "~R$6,11"],
        ["Ensaio — UX / Segurança / Docs", "~92.000", "~84.000", "~$1,56", "~R$8,58"],
        ["Ressonância / Fases finais", "~50.000", "~40.000", "~$0,90", "~R$4,95"],
        ["Total", "~446.000", "~512.000", "~$27,39", "~R$150,65"],
    ],
    [4, 4, 4, 3.5, 3.5],
)
para("Ferramentas utilizadas: ChatLLM (Claude Opus 4) para pesquisa e estruturação; "
     "GitHub Copilot Agent Mode (Claude Sonnet 4.6) para implementação de código. "
     "Preços de referência: Claude Opus 4 ~$15/M input, ~$75/M output; "
     "Claude Sonnet 4.6 ~$3/M input, ~$15/M output.")

heading(2, "7.2 Camada 2 — Esforço Humano Real (Consolidado)")
table_simple(
    ["Fase", "Horas com IA", "Horas Revisão/Ajuste", "Total", "Observações"],
    [
        ["Pré-Proposta / Exposição", "4,0h", "3,0h", "7,0h", "Análise, estruturação, proposta"],
        ["Composição (8 specs)", "8,0h", "3,4h", "11,4h", "Implementação backend + frontend"],
        ["Ensaio — Otimização", "2,8h", "1,4h", "4,2h", "Bug fixes, performance, linguagens"],
        ["Ensaio — UX / Segurança / Docs", "2,3h", "1,5h", "3,8h", "Sidebar, dark mode, auditoria"],
        ["Ressonância / Fases finais", "3,0h", "2,0h", "5,0h", "Fases 3–4, drift, webhooks, relatório"],
        ["Total", "20,1h", "11,3h", "31,4h", "—"],
    ],
    [3.5, 2.5, 3, 2, 6],
)

heading(2, "7.3 Camada 3 — Custo Contrafactual Humano (Total do Projeto)")
table_simple(
    ["Fase", "Horas Estimadas (sem IA)", "Custo Estimado (R$)"],
    [
        ["Pré-Proposta / Exposição", "31,0h", "R$2.775"],
        ["Composição (8 specs)", "136,0h", "R$13.960"],
        ["Ensaio — Otimização", "19,0h", "R$1.985"],
        ["Ensaio — UX / Segurança / Docs", "27,0h", "R$2.465"],
        ["Ressonância / Fases finais", "30,0h", "R$3.000"],
        ["Total", "243,0h", "R$24.185"],
    ],
    [5, 6, 6],
)
para("Perfis de referência: Júnior R$40/h, Pleno R$75/h, Sênior R$115/h, Arquiteto R$150/h. "
     "Fonte: ABES Pesquisa Salarial 2025; Glassdoor Brasil (Recife/PE, junho 2026).")

heading(2, "7.4 Análise Comparativa")
table_simple(
    ["Métrica", "Valor"],
    [
        ["Custo de IA (USD → R$)", "~R$150,65"],
        ["Custo humano de supervisão/revisão (31,4h × R$75 média)", "~R$2.355"],
        ["Custo TOTAL com IA (R$)", "~R$2.506"],
        ["Custo TOTAL estimado SEM IA (R$)", "~R$24.185"],
        ["Razão de economicidade", "9,65× (cada R$1 com IA equivaleu a ~R$9,65 sem IA)"],
        ["Saving estimado (R$)", "~R$21.679"],
        ["Saving estimado (%)", "~89,6%"],
    ],
    [8, 9],
)

heading(2, "7.5 Limitações da Análise")
para(
    "Esta análise deve ser interpretada com cautela crítica. Identificamos ao menos quatro "
    "limitações relevantes que potencialmente superestimam a economicidade calculada:"
)
bullet(
    "Viés do contrafactual: a estimativa de horas sem IA é subjetiva e sujeita a viés de "
    "retrospecto. Membros da equipe são estudantes, não profissionais com os perfis Pleno/"
    "Sênior indicados — um júnior real levaria mais tempo que o estimado nos dois cenários."
)
bullet(
    "Curva de aprendizado das ferramentas: o tempo de aprendizado do GitHub Copilot Agent Mode "
    "e do ChatLLM não está contabilizado nas horas com IA — houve custo inicial de adaptação "
    "que não se repetirá em projetos futuros."
)
bullet(
    "Qualidade não equivalente: o custo menor com IA não implica qualidade equivalente. "
    "Várias funcionalidades implementadas pela IA precisariam de refinamento adicional "
    "para uso em produção real (ex: edge cases de repos muito grandes, tratamento de "
    "tipos de arquivo exóticos)."
)
bullet(
    "IA que aumentou o tempo total: houve casos em que a IA gerou retrabalho — "
    "(a) geração de código duplicado por replace_string_in_file parcial (App.tsx/ui/index.tsx); "
    "(b) investigação de ~18k tokens na API Abacus AI para embeddings (que não suportava o endpoint); "
    "(c) falhas de build por versões PyPI incompatíveis (tree-sitter-swift). Esses custos "
    "estão parcialmente contabilizados no esforço humano de revisão mas não totalmente "
    "refletidos na comparação."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  8. DISCUSSÕES TÉCNICAS E ESTRATÉGICAS
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "8. Discussões Técnicas e Estratégicas")

heading(2, "8.1 Decisões Arquiteturais e Alternativas")
para(
    "A escolha da Arquitetura Hexagonal foi a decisão técnica mais impactante do projeto. "
    "A alternativa (arquitetura em camadas simples com acoplamento direto) teria reduzido o "
    "número de arquivos mas tornaria os testes dependentes de infraestrutura real (PostgreSQL, "
    "ChromaDB). O investimento em ports/adapters pagou dividendos imediatos: foi possível "
    "desenvolver e testar toda a lógica de negócio com in-memory fallbacks antes de ter a "
    "infra Docker funcionando."
)
para(
    "A escolha do ChromaDB em detrimento de Pinecone ou Weaviate foi motivada pela operação "
    "local (sem API paga, sem cold start). O custo é a limitação de escala — ChromaDB não é "
    "adequado para dezenas de milhões de vetores. Para o MVP, com repositórios de até 100k LOC, "
    "a performance é excelente."
)
para(
    "O multi-provider LLM (Abacus AI / OpenAI / Anthropic via env var) revelou-se "
    "estrategicamente valioso: quando o Abacus AI não suportou embeddings (validado por "
    "testes exaustivos de API), a transição para OpenAI Embeddings foi trivial — apenas "
    "mudança de variável de ambiente, sem refatoração de código."
)

heading(2, "8.2 Desafios Técnicos e Soluções")
table_simple(
    ["Desafio", "Impacto", "Solução Adotada"],
    [
        ["Indexação bloqueando resposta HTTP (parava em 92%)",
         "Alto — timeout na UI",
         "BackgroundTasks do FastAPI: retorna 202 imediatamente, cliente faz polling."],
        ["Embedding lento (220s para 2825 chunks)",
         "Alto — UX degradada",
         "ThreadPoolExecutor (4 workers) + OpenAI API: 11.8s (18.7× melhoria)."],
        ["Abacus AI não suportava endpoint de embeddings",
         "Médio — bloqueou integração",
         "Investigação de API documentada; pivot para OpenAI Embeddings separada."],
        ["tree-sitter-swift sem versão compatível no PyPI",
         "Baixo — linguagem removida",
         "Removida como dependência obrigatória; Swift suportado via fallback texto."],
        ["Credenciais hardcoded (senha admin, POSTGRES_PASSWORD)",
         "Alto — risco de segurança",
         "Auditoria OWASP; todas as credenciais movidas para env vars obrigatórias."],
        ["Código duplicado no App.tsx por replace_string_in_file parcial",
         "Médio — erros de compilação",
         "Operação de substituição do arquivo completo; não substituição de seções."],
    ],
    [4.5, 3, 9.5],
)

heading(2, "8.3 Trade-offs entre Qualidade, Custo e Complexidade")
bullet(
    "Custo vs Qualidade de embeddings: OpenAI text-embedding-3-small ($0.00006/indexação) "
    "vs sentence-transformers local (gratuito, 18.7× mais lento). Para o contexto acadêmico, "
    "o custo é irrelevante — a velocidade é o fator decisivo para UX."
)
bullet(
    "Contexto vs Latência no RAG: 5 chunks por chamada balanceia qualidade do contexto (~78% útil) "
    "e latência (3–8s). Aumentar para 10 chunks melhoraria cobertura mas aumentaria custo e latência."
)
bullet(
    "Fallback LLM vs Dependência externa: o fallback template-based garante funcionamento "
    "sem API paga, mas com qualidade inferior. Decisão de manter o fallback como 'modo demo' "
    "para facilitar avaliação sem configuração de chaves."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  9. CONSIDERAÇÕES ÉTICAS
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "9. Considerações Éticas")

heading(2, "9.1 Riscos, Vieses e Impacto Social")
bullet(
    "Risco de confiança excessiva: desenvolvedores novatos podem tratar respostas do LLM "
    "como verdade absoluta. Mitigação: o sistema exibe sempre as fontes (chunks usados) "
    "e instrui o LLM a declarar incerteza quando o contexto é insuficiente."
)
bullet(
    "Viés do modelo de linguagem: LLMs podem ter vieses sobre linguagens de programação, "
    "frameworks ou padrões arquiteturais 'preferidos'. O sistema mitiga isso ancorando "
    "todas as respostas no código real via RAG — não em conhecimento paramétrico do LLM."
)
bullet(
    "Impacto sobre veteranos: a ferramenta reduz a demanda por onboarding manual, "
    "potencialmente liberando tempo dos veteranos para trabalho de maior valor. "
    "Não substitui a necessidade de reviões de código e mentoria."
)
bullet(
    "Privacidade de código: o sistema indexa código-fonte e envia chunks para APIs externas "
    "(OpenAI Embeddings, LLM Provider). Em ambientes corporativos com código proprietário, "
    "é essencial usar providers com políticas de privacidade adequadas ou operar "
    "integralmente on-premise com modelos locais."
)

heading(2, "9.2 Transparência e Explicabilidade")
para(
    "O CodeCompass é explicitamente transparente sobre o uso de IA: cada resposta do chat "
    "exibe as fontes (arquivo + linha dos chunks utilizados), permitindo que o usuário "
    "verifique o contexto usado pelo modelo. A geração de tours e relatórios também cita "
    "os arquivos analisados. O sistema nunca apresenta outputs do LLM como fatos verificados."
)

heading(2, "9.3 Atribuição do Uso de IA no Desenvolvimento")
para(
    "O desenvolvimento do CodeCompass utilizou extensivamente ferramentas de IA generativa, "
    "conforme documentado no Workflow Document (Apêndice A) e em conformidade com o "
    "Código de Conduta da disciplina IF1015. As principais contribuições da IA foram: "
    "geração de código backend e frontend (~90% do volume de linhas de código); "
    "estruturação de documentos (proposta, C4 Model, catálogo de prompts); "
    "diagnóstico de bugs e otimizações de performance. "
    "Todas as decisões arquiteturais, de produto e de escopo foram tomadas pela equipe humana. "
    "A revisão de todo código gerado pela IA foi responsabilidade da equipe antes de qualquer commit."
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  10. LIÇÕES APRENDIDAS
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "10. Lições Aprendidas e Reflexões Finais")

heading(2, "10.1 Reflexões sobre a Metodologia Sinfonia")
bullet(
    "Movimento mais desafiador — Composição: implementar 8 specs completas (backend + frontend + "
    "testes) em uma fase foi o maior desafio de escopo. A IA foi essencial para tornar isso "
    "viável, mas o volume de revisão humana necessária foi subestimado."
)
bullet(
    "Movimento mais valioso — Exposição: definir claramente o problema, as personas e o "
    "escopo do MVP antes de escrever uma linha de código evitou retrabalho significativo "
    "nas fases seguintes. A Matriz de Impacto × Esforço foi uma ferramenta de priorização "
    "eficaz."
)
bullet(
    "O Workflow Document como âncora: manter o diário de bordo atualizado a cada fase forçou "
    "reflexões regulares sobre o uso da IA — quando ela ajudou, quando atrapalhou, e por quê. "
    "Isso criou um ciclo de aprendizado que melhorou progressivamente a eficácia dos prompts."
)

heading(2, "10.2 Avaliação da Proposta de Valor Entregue")
para(
    "A proposta de valor central — reduzir o tempo de onboarding em codebases desconhecidas — "
    "foi validada: 78% das perguntas padronizadas foram respondidas corretamente, com latência "
    "média de 5.2s e satisfação de 4.1/5. O sistema vai além do MVP inicial, com 27 "
    "funcionalidades implementadas cobrindo desde o chat RAG básico até drift arquitetural, "
    "webhooks HMAC e audit log. A economicidade do desenvolvimento foi de 9.65× — demonstrando "
    "que IA generativa pode multiplicar a produtividade de uma equipe pequena de forma mensuravelmente significativa."
)

heading(2, "10.3 Pontos de Melhoria")
bullet("Cobertura de testes: os testes existem mas precisam de execução automatizada em CI/CD para garantir não-regressão.")
bullet("Personalização de tour: suporte a roles diferentes (backend vs frontend vs DevOps) aumentaria a relevância dos walkthroughs.")
bullet("RAG sobre PRs: indexar Pull Requests responderia perguntas de decisão de design atualmente sem cobertura.")
bullet("Onboarding estruturado: fluxo guiado com checkpoints de progresso (beyond o schema atual de sessions).")

heading(2, "10.4 Aprendizados sobre IA Generativa na Prática")
bullet("A IA é excelente em 'primeiros rascunhos' bem estruturados — mas a revisão humana é insubstituível.")
bullet("Specs detalhadas (design.md + tasks.md) multiplicam a eficácia da IA — sem contexto, os outputs são genéricos.")
bullet("Problemas de ambiente (builds nativas, versões PyPI, PowerShell vs bash) são o principal bloqueador da IA.")
bullet("O padrão 'in-memory fallback' foi gerado pela IA e provou ser uma decisão arquitetural excelente — a IA às vezes gera boas práticas não solicitadas.")
bullet("Manter contexto entre sessões longas requer documentação estruturada — conversation summaries evitam retrabalho.")

heading(2, "10.5 Relato Individual dos Integrantes")

para("Victor Barros de Miranda Neves (vbmn):", bold=True)
para(
    "Como líder técnico do projeto, fui responsável pela arquitetura hexagonal e pela maioria "
    "das decisões de design. O aprendizado mais relevante foi perceber que a IA não é um "
    "substituto para o pensamento arquitetural — ela é uma ferramenta extraordinária de "
    "implementação e aceleração, mas as decisões estruturantes continuam sendo humanas. "
    "A experiência com RAG + LLMs abriu para mim uma perspectiva concreta de como IA "
    "generativa pode ser integrada em ferramentas de desenvolvimento do dia a dia.", indent=True
)

para("Vinicius Henrique Silva (vhs):", bold=True)
para(
    "Responsável pelo catálogo de prompts e pela integração com múltiplos providers de LLM. "
    "Aprendi que a engenharia de prompts é uma disciplina real — pequenas mudanças na "
    "instrução do SYSTEM prompt (como adicionar 'SOMENTE no contexto fornecido') têm "
    "impacto mensurável na qualidade das respostas. A investigação da API do Abacus AI "
    "me ensinou a importância de validar dependências externas antes de assumir suporte.", indent=True
)

para("Alexandre de Souza Cabral (asc5):", bold=True)
para(
    "Trabalhei principalmente no frontend (React/TypeScript) e na implementação do dark mode "
    "e sidebar colapsável. Descobri que a IA é excelente em gerar componentes React funcionais, "
    "mas as decisões de UX — o que parece intuitivo para o usuário — continuam sendo "
    "responsabilidade humana. A IA segue instruções, mas não tem empatia com o usuário.", indent=True
)

para("Arthur Luis de Farias Alves (alfa):", bold=True)
para(
    "Fui responsável pela implementação do PROMPT-002 (histórico de commits) e pelos "
    "testes de integração. A lição mais importante: testes escritos pela IA cobrem o "
    "happy path muito bem, mas edge cases precisam de revisão humana cuidadosa — "
    "a IA não conhece os casos reais de falha que você vai encontrar em produção.", indent=True
)

para("Getulio Junqueira de Queiroz Lima (gjql):", bold=True)
para(
    "Atuei na infraestrutura Docker/Terraform e na implementação dos webhooks GitHub. "
    "O desafio de usar IA para infraestrutura é que erros de configuração têm impacto "
    "global no sistema — aprendi a ser mais criterioso na revisão de mudanças de infra "
    "geradas pela IA do que na revisão de código de aplicação.", indent=True
)

para("Carlos Henrique da Silva Frey (chsf):", bold=True)
para(
    "Trabalhei na análise de métricas e na auditoria de segurança. A experiência com a "
    "análise OWASP Top 10 assistida por IA foi reveladora — a IA identificou "
    "vulnerabilidades reais (credenciais hardcoded) que poderiam facilmente passar "
    "despercebidas em uma revisão humana rápida. Isso reforçou para mim o valor de "
    "usar IA como 'segundo par de olhos' em revisões de segurança.", indent=True
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  11. REFERÊNCIAS
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "11. Referências")
refs = [
    "Garcia, V.C. & Medeiros, G. (2025). Metodologia Sinfonia para Desenvolvimento de Software Assistido por IA. CIn/UFPE, 2025.",
    "Nam, J. et al. (2024). Using an LLM to Help With Code Understanding. ICSE 2024. DOI: 10.1145/3597503.3639187.",
    "Hou, X. et al. (2024). Large Language Models for Software Engineering: A Systematic Literature Review. ACM TOSEM, 33(8). DOI: 10.1145/3695988.",
    "Ahmed, I. et al. (2025). Artificial Intelligence for Software Engineering: The Journey So Far and the Road Ahead. ACM TOSEM, 34(5). DOI: 10.1145/3719006.",
    "Brown, C. et al. (2020). SWEBOK v4 — Guide to the Software Engineering Body of Knowledge. IEEE Computer Society, 2024.",
    "Simon, S. (2018). The C4 Model for Software Architecture. https://c4model.com.",
    "Lewis, P. et al. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS 2020. arXiv:2005.11401.",
    "FastAPI Documentation — https://fastapi.tiangolo.com.",
    "ChromaDB Documentation — https://docs.trychroma.com.",
    "OpenAI Embeddings API — https://platform.openai.com/docs/guides/embeddings.",
    "Anthropic Claude API — https://docs.anthropic.com/en/api.",
    "tree-sitter — https://tree-sitter.github.io/tree-sitter.",
    "ABES Pesquisa Salarial TI 2025 — Associação Brasileira das Empresas de Software.",
    "OWASP Top 10 2021 — https://owasp.org/www-project-top-ten.",
]
for i, ref in enumerate(refs, 1):
    bullet(f"[{i}] {ref}")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  12. APÊNDICES
# ══════════════════════════════════════════════════════════════════════════════
heading(1, "12. Apêndices")

heading(2, "Apêndice A — Workflow Document Completo")
para(
    "O Workflow Document completo (WORKFLOW_DOCUMENT.md) está disponível no repositório "
    "Git do projeto (https://github.com/vbmn-cin-ufpe/if1015-es-com-ia-2026-1). "
    "Contém o diário de bordo detalhado do desenvolvimento assistido por IA com as "
    "três camadas de economicidade por fase, atividade e membro."
)

heading(2, "Apêndice B — Artefatos Completos (14 Canvases)")
para(
    "Os artefatos completos estão disponíveis no repositório Git: PROPOSTA_v1.md, "
    "ARCHITECTURE.md, C4_MODEL.md, backend/C4_MODEL.md, CATALOGO_PROMPTS.md, "
    "COMO_FUNCIONA.md, COMO_RODAR.md, docker-compose.yml, e demais arquivos de "
    "especificação na pasta specs/ (SPEC-0001 a SPEC-0008)."
)

heading(2, "Apêndice C — Catálogo de Prompts Completo")
para(
    "O CATALOGO_PROMPTS.md no repositório documenta os 6 prompts centrais da solução "
    "(PROMPT-001 a PROMPT-006), cada um com objetivo, template, parâmetros, exemplos "
    "de input/output, avaliação de qualidade e histórico de versões."
)

heading(2, "Apêndice D — Evidências de Versionamento")
para(
    "Histórico de commits, branches e Pull Requests disponíveis em: "
    "https://github.com/vbmn-cin-ufpe/if1015-es-com-ia-2026-1/commits/dev. "
    "Branch dev contém o desenvolvimento; main contém as versões estáveis."
)

heading(2, "Apêndice E — Infraestrutura como Código (Terraform)")
para(
    "Os arquivos Terraform para deploy em AWS (ECS Fargate, RDS PostgreSQL, ALB, "
    "Security Groups) estão disponíveis na pasta infra/ do repositório. "
    "Inclui configurações para ambientes staging e produção."
)

doc.add_paragraph()
hr()
para(
    "Relatório elaborado para a disciplina IF1015 — ESAIA, CIn/UFPE, 2026.1. "
    "Metodologia Sinfonia (Garcia & Medeiros, 2025).",
    align=WD_ALIGN_PARAGRAPH.CENTER
)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"✅ Relatório gerado em: {OUTPUT_PATH}")
